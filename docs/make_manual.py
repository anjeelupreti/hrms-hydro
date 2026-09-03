"""Build the Xenex HRMS user manual as a .docx.

Written as code rather than typed in Word for one reason: the manual has to be
rebuilt whenever the product moves, and a document somebody re-typed is a
document that drifts. The figures come from a live capture run, the credentials
from the seed, and the rules from the modules that enforce them.

    python docs/make_manual.py <figures-dir> <output.docx>

**The table of contents is a field, not a list.** python-docx cannot compute
page numbers — that is the layout engine's job — so a `TOC` field is written
and Word fills it in. The document opens with "update this field" instructions
next to it, because a field nobody refreshes shows nothing at all, which looks
like a broken document rather than an unrefreshed one.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x3F, 0x37, 0xC9)
MUTED = RGBColor(0x55, 0x5C, 0x6B)
ALARM = RGBColor(0xB3, 0x1F, 0x32)

FIGURES = Path(".")


# ── Small helpers over python-docx ───────────────────────────────────────


def _shade(cell, hex_colour: str) -> None:
    """Fill a table cell. python-docx has no API for this, so it is XML."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def _field(paragraph, instruction: str) -> None:
    """Insert a Word field — used for the TOC and for page numbers."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field”."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, placeholder, end):
        run._r.append(node)


class Manual:
    def __init__(self, figures: Path):
        self.figures = figures
        self.doc = Document()
        self._styles()

    def _styles(self) -> None:
        base = self.doc.styles["Normal"]
        base.font.name = "Calibri"
        base.font.size = Pt(10.5)
        base.paragraph_format.space_after = Pt(6)
        base.paragraph_format.line_spacing = 1.15
        for name, size, colour in (
            ("Heading 1", 20, ACCENT),
            ("Heading 2", 15, ACCENT),
            ("Heading 3", 12, RGBColor(0x1A, 0x1D, 0x26)),
        ):
            style = self.doc.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = colour
            style.font.bold = True

    # ── Content primitives ───────────────────────────────────────────────

    def h1(self, text: str, page_break: bool = True):
        if page_break:
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return self.doc.add_heading(text, level=1)

    def h2(self, text: str):
        return self.doc.add_heading(text, level=2)

    def h3(self, text: str):
        return self.doc.add_heading(text, level=3)

    def p(self, text: str = "", bold: bool = False, colour: RGBColor | None = None):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        if colour is not None:
            run.font.color.rgb = colour
        return paragraph

    def rich(self, *chunks):
        """A paragraph of (text, bold) pairs, for a sentence with emphasis in it."""
        paragraph = self.doc.add_paragraph()
        for text, bold in chunks:
            run = paragraph.add_run(text)
            run.bold = bold
        return paragraph

    def bullets(self, items) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def steps(self, items) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Number")

    def note(self, text: str, kind: str = "note") -> None:
        """A one-cell shaded table — the only reliable call-out box in Word."""
        fill, colour = {
            "note": ("EEF0FF", ACCENT),
            "warn": ("FFF3E0", RGBColor(0x9A, 0x5B, 0x00)),
            "stop": ("FDECEE", ALARM),
        }[kind]
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        _shade(cell, fill)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = colour
        self.doc.add_paragraph()

    def table(self, headers, rows, widths=None) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            _shade(cell, "3F37C9")
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                run = cells[index].paragraphs[0].add_run(str(value))
                run.font.size = Pt(9.5)
        if widths:
            for row in table.rows:
                for index, width in enumerate(widths):
                    row.cells[index].width = Inches(width)
        self.doc.add_paragraph()

    def figure(self, name: str, caption: str, width: float = 6.3) -> None:
        path = self.figures / f"{name}.png"
        if not path.exists():
            # A missing figure is noted, not silently skipped: the reader should
            # know a picture was meant to be here.
            self.p(f"[figure missing: {name}]", colour=MUTED)
            return
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(width))
        caption_paragraph = self.doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.add_run(caption)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        run.italic = True

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        self.doc.save(str(path))
        return path


# ── The document ─────────────────────────────────────────────────────────


def cover(m: Manual) -> None:
    for _ in range(5):
        m.doc.add_paragraph()
    title = m.doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Xenex HRMS")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    sub = m.doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("User Manual and Functional Specification")
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    m.doc.add_paragraph()
    org = m.doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run("Vision Lumbini Urja Company Limited")
    run.font.size = Pt(13)
    run.font.bold = True

    seat = m.doc.add_paragraph()
    seat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = seat.add_run(
        "Butwal-8, Rupandehi · Lumbini Province\n"
        "Seti Nadi Hydropower · Sanjen Jalavidyut · Marsyangdi Corridor Transmission"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    for _ in range(8):
        m.doc.add_paragraph()
    note = m.doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This document describes the system as installed. Every screenshot is a "
        "live capture of the running application, and every rule stated here is "
        "one the software enforces."
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MUTED


def contents(m: Manual) -> None:
    m.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    m.doc.add_heading("Contents", level=1)
    m.note(
        "This table is generated by Word. If it shows a placeholder instead of "
        "page numbers, click it once, then press F9 (or right-click, then Update "
        "Field, then Update entire table). Word cannot compute page numbers until "
        "it has laid the document out, so no tool can fill them in beforehand.",
        "note",
    )
    _field(m.doc.add_paragraph(), r'TOC \o "1-3" \h \z \u')


def introduction(m: Manual) -> None:
    m.h1("1. About this system")
    m.rich(
        ("Xenex HRMS is a human-resources system installed for ", False),
        ("Vision Lumbini Urja Company Limited", True),
        (" and the three ventures it holds. It is a single-company installation: "
         "there is no sign-up page, no tenant switching, and no marketing site. "
         "Everybody who uses it has an account created for them from inside the "
         "product.", False),
    )
    m.h2("1.1 The idea the system is built on")
    m.rich(
        ("One record per person", True),
        (". Payroll, attendance, leave, training, expenses and the approval "
         "chains all read the same employee row. Nothing is copied between "
         "modules, so nothing can disagree — an address corrected once is "
         "corrected everywhere it appears.", False),
    )
    m.h2("1.2 The group")
    m.table(
        ["Code", "Entity", "Kind", "Status"],
        [
            ["VLUCL", "Vision Lumbini Urja Company Limited", "Holding company", "Butwal, Rupandehi"],
            ["SNHL", "Seti Nadi Hydropower Limited", "Project company", "25.0 MW, in operation"],
            ["SJCL", "Sanjen Jalavidyut Company Limited", "Project company", "42.5 MW, under construction"],
            ["MCTL", "Marsyangdi Corridor Transmission Limited", "Subsidiary", "Licensed"],
        ],
        widths=[0.8, 2.6, 1.4, 1.7],
    )
    m.p(
        "An employee belongs to one primary company and may be seconded to any "
        "number of others. Company codes are not decoration: they form part of "
        "every memorandum number the system issues."
    )
    m.h2("1.3 Built for how payroll is actually run here")
    m.bullets([
        "Bikram Sambat alongside the Gregorian calendar, throughout",
        "SSF and Provident Fund as alternatives, never both at once",
        "CIT, gratuity, and the retirement-relief ceilings",
        "Nepali income-tax slabs, by taxpayer type and fiscal year",
        "TDS, computed from the same slabs",
    ])


def access(m: Manual) -> None:
    m.h1("2. Getting in")
    m.h2("2.1 Where the system lives")
    m.table(
        ["What", "Address"],
        [
            ["The application", "http://localhost:3001"],
            ["Backend API (JSON only)", "http://localhost:8001"],
            ["Django administration", "http://localhost:8001/admin/"],
            ["Outgoing mail (Mailhog)", "http://localhost:8026"],
        ],
        widths=[2.4, 4.1],
    )
    m.note(
        "No mail leaves the machine in this installation. Suspension notices, "
        "memorandum alerts and password resets all land in the Mailhog inbox. "
        "Real delivery has to be configured under Settings, Email before "
        "anything reaches a real address.",
        "warn",
    )

    m.h2("2.2 Signing in")
    m.figure("fig-login", "Figure 2.1 — The sign-in screen")
    m.steps([
        "Open http://localhost:3001 in a browser. You are sent to the sign-in screen.",
        "Type your username. It is your first name and last name joined by a full stop, for example sushma.ghimire.",
        "Type your password. Click the eye at the right of the box to check what you have typed; click it again to hide it.",
        "Click Sign in.",
    ])
    m.p(
        "The Sign in button stays greyed for a moment after the page appears. "
        "That is deliberate: the button is only enabled once the page is fully "
        "loaded, so that pressing Enter early cannot submit your password in a "
        "way that would put it in the browser's address bar."
    )
    m.h3("If you cannot get in")
    m.table(
        ["What you see", "What it means"],
        [
            ["No active account found with the given credentials", "The username or password is wrong."],
            ["This account is suspended until …", "A suspension is in force. Speak to HR; the date is on the message."],
            ["A code is asked for", "Two-factor authentication is on for your account. Enter the code from your authenticator app, or a backup code."],
        ],
        widths=[2.6, 3.9],
    )
    m.note(
        "There is no self-service sign-up and no Google or Microsoft sign-in. "
        "Accounts are created inside the product by HR, which is why the "
        "sign-in screen says so.",
        "note",
    )

    m.h2("2.3 The first password")
    m.p(
        "An account created for you arrives with a temporary password by email. "
        "The first time you sign in, the system asks you to choose your own "
        "before it lets you go any further. That screen cannot be skipped: a "
        "password somebody else has seen is not private."
    )


def roles(m: Manual) -> None:
    m.h1("3. Who can do what")
    m.p(
        "Four roles. The difference between them is the single most important "
        "thing to understand about this system, because it changes what appears "
        "on screen as well as what is permitted."
    )

    m.h2("3.1 The four roles")
    m.table(
        ["Role", "Held by", "What it is for"],
        [
            ["Owner", "One account", "The account the system was installed under. Appoints HR admins. Also reaches the Django administration site."],
            ["HR admin", "Senior HR staff", "Runs the system. Creates, edits and deletes. Everything except appointing another admin."],
            ["HR officer", "HR staff", "Operates the system. Reads and edits. Cannot create or delete anything."],
            ["Employee", "Everybody else", "Their own record, their own requests, and whatever a workflow has named them in."],
        ],
        widths=[1.0, 1.4, 4.1],
    )

    m.h2("3.2 The rule that separates an admin from an officer")
    m.rich(
        ("An HR admin creates. An HR officer operates.", True),
        (" An officer may open any employee and change their address, their "
         "phone number, their manager, their department. An officer may not add "
         "a new employee, and may not delete anything at all. The refusal names "
         "the verb: “This account may not create records of this kind.”", False),
    )
    m.p(
        "This is not a list of forbidden buttons; it is one rule applied to "
        "every kind of record in the system. If an officer can edit a thing, "
        "they cannot create or delete that same thing."
    )
    m.note(
        "Deletion never widens. An HR officer promoted to admin gains the "
        "ability to delete; nobody gains it by any other route, because "
        "deletion is the one act with no undo.",
        "warn",
    )

    m.h2("3.3 What an officer holds by default")
    m.p("Appointing somebody an HR officer gives them, without any further setup:")
    m.bullets([
        "people.view and people.manage — the employee record",
        "attendance.manage — logs, shifts, corrections",
        "leave.approve — leave requests",
        "workplace.manage — assets, training, helpdesk, timesheets, remote work",
        "reports.view and dashboard.view",
    ])
    m.p("Deliberately not included, and granted one person at a time if needed:")
    m.bullets([
        "settings.manage — salary components, tax slabs, the memorandum vocabulary. Shaping the system is the admin's.",
        "payroll.view and payroll.run — money",
        "expenses.manage, recruitment.manage, crm.manage",
        "mail.access — the company mailbox",
        "people.admin — never grantable to anybody, because a grantable “grant permissions” is how an officer becomes an owner in two steps",
    ])

    m.h2("3.4 Scope: being somebody's manager")
    m.p(
        "Scope is separate from permission, and getting the two confused is the "
        "usual mistake. Being named as somebody's manager gives you reach over "
        "their records — their leave requests arrive with you — without granting "
        "you any capability. The reverse also holds: holding leave.approve with "
        "no reports approves nothing."
    )

    m.h2("3.5 Where roles are set")
    m.p(
        "Sidebar, People, Roles and permissions. The owner may appoint anybody "
        "to anything. An HR admin may promote and demote officers and employees, "
        "and may not create or remove another admin — that stays with the owner, "
        "because an admin who can create admins is an owner by a longer route."
    )
    m.figure("fig-team", "Figure 3.1 — Roles and permissions")


def employees(m: Manual) -> None:
    m.h1("4. Employees")
    m.p(
        "The employee record is the centre of the system. Every other module "
        "reads it, so this chapter is worth reading before the others."
    )

    m.h2("4.1 The directory")
    m.figure(
        "fig-employees",
        "Figure 4.1 — The employee list, with the arrangement every list in the system copies",
    )
    m.p("Numbered on the figure above:")
    m.table(
        ["#", "What it is", "What it does"],
        [
            ["1", "Search", "Matches name, employee code and email. Sits under the title, on the left of the filter card — the same place on every list in the system."],
            ["2", "Current / Past employees", "Splits the roster from leavers. Picking a specific status below overrides this."],
            ["3", "Status chips", "Both the count and the filter. The number is the whole directory, not the page you are looking at. Click one to narrow the list."],
            ["4", "A row", "Click anywhere on it to open that person."],
            ["5", "Add Employee", "HR admin and owner only. An officer does not get a working button here."],
        ],
        widths=[0.35, 1.5, 4.65],
    )
    m.h3("Step by step: finding somebody")
    m.steps([
        "Click Employees in the sidebar, under People.",
        "Type part of their name or their employee code into the search box.",
        "If you know their status, click the matching chip: Active, On leave, Suspended, Resigned or Terminated.",
        "To narrow further, use the Department, Designation or Company dropdowns beside the search box.",
        "Click the row to open the profile.",
    ])
    m.note(
        "Pagination sits in its own bar at the foot of the list and disappears "
        "entirely when everything fits on one page. A control that can only say "
        "1 of 1 is furniture.",
        "note",
    )

    m.h2("4.2 One profile, reached from everywhere")
    m.figure("fig-employee-profile", "Figure 4.2 — An employee profile")
    m.p(
        "There is one employee profile page. A click on a name in payroll, in a "
        "leave request, in an approval chain or in the org chart all arrive at "
        "the same address. The tabs across the top are:"
    )
    m.table(
        ["Tab", "What is on it"],
        [
            ["Overview", "The card, the current position, the manager, at-a-glance figures"],
            ["Record", "Employment record: department, designation, dates, companies"],
            ["Payroll", "Salary structure and past payslips"],
            ["Attendance", "Logs, lateness, the monthly pattern"],
            ["Training", "Programmes requested, enrolled on, completed"],
            ["Projects", "What they are assigned to"],
            ["Personal", "Contact details, addresses, emergency contacts, dependants, nominees, education"],
            ["Conduct", "Awards, disciplinary actions, suspensions"],
            ["Lifecycle", "Promotions, transfers, resignation, termination"],
            ["Activity", "Who changed what on this record, and when"],
        ],
        widths=[1.1, 5.4],
    )

    m.h2("4.3 Editing somebody")
    m.figure("fig-employee-edit", "Figure 4.3 — The employee form")
    m.h3("Step by step")
    m.steps([
        "Open the employee.",
        "Click Edit at the top right of the profile card. The form opens as a dialog over the page.",
        "Change what you need. The form is grouped: identity first, then position, then companies, then contact, then bank and statutory details.",
        "Click Save. The dialog closes and the profile behind it updates.",
    ])
    m.h3("The fields, and what they are for")
    m.table(
        ["Field", "Notes"],
        [
            ["Reports to", "This field is the organisation chart. Everything drawn on the Org chart page comes from it, as does every team list and manager notification. A company where nobody has set it has an empty chart."],
            ["Corporate post", "What somebody is appointed to, and what their grade follows. Deputy Manager, Senior Engineer."],
            ["Corporate role", "What they are responsible for. Two Deputy Managers hold different roles, and somebody promoted usually keeps running the same site. That is why it is a second field rather than the same one."],
            ["Primary company", "Which of the four entities employs them."],
            ["Secondary companies", "Where else they work. A multi-select: add and remove freely."],
            ["Probation ends", "Blank means not on probation."],
            ["Employment status", "Active, On leave, Resigned, Terminated. Suspended is deliberately absent, for the reason in 4.5."],
            ["Office cell and office email", "Correspondence details. Visible to colleagues, because that is what a directory is for."],
            ["Personal cell and personal email", "Private. Visible only to HR and to the person themselves."],
            ["Permanent and temporary address", "Private, as above."],
            ["Blood group", "Visible to colleagues. It is on the ID card for the same reason it is here."],
            ["Bank, citizenship, PAN, SSF, PF, CIT", "Private. HR and the person only."],
        ],
        widths=[1.6, 4.9],
    )
    m.note(
        "What a colleague sees is decided when the record is served, not by "
        "hiding fields on screen. An ordinary employee opening a co-worker's "
        "profile is not sent the private fields at all.",
        "note",
    )

    m.h2("4.4 The organisation chart")
    m.figure("fig-org-chart", "Figure 4.4 — The organisation chart")
    m.rich(
        ("The chart has one source: the ", False),
        ("Reports to", True),
        (" field on each employee. It is drawn by walking that field upwards. "
         "There is no separate hierarchy to maintain and no second place to "
         "correct. If somebody is in the wrong place on the chart, open them and "
         "change who they report to.", False),
    )

    m.h2("4.5 Suspension")
    m.figure("fig-employee-conduct", "Figure 4.5 — The Conduct tab")
    m.rich(
        ("Suspended is not a status you can choose from the dropdown.", True),
        (" It comes from a suspension record, so that the interval, the reason "
         "and the account lock always move together. A dropdown would let "
         "somebody be marked suspended with no record of since when, and still "
         "able to sign in.", False),
    )
    m.h3("Step by step: suspending somebody")
    m.steps([
        "Open the employee and click the Conduct tab.",
        "Click Suspend.",
        "Enter the start date. Leave the end date blank for an open-ended suspension pending an inquiry, or set one for a fixed period.",
        "Write the reason. It appears on the banner and in the notice sent to them.",
        "Click Suspend.",
    ])
    m.p("Four things happen at once, and they cannot happen separately:")
    m.bullets([
        "The roster status becomes Suspended, and a banner appears on the profile",
        "The account is locked, so they cannot sign in from that moment",
        "The sign-in screen tells them they are suspended, and until when, rather than saying the password is wrong",
        "A notice is emailed to them, to their personal address if the record has one",
    ])
    m.h3("Step by step: lifting a suspension")
    m.steps([
        "Open the employee, Conduct tab.",
        "Click Lift on the active suspension.",
        "Choose the outcome. This is required, and it is not a formality: the three answers lead to different places.",
        "Add a note, and click Lift.",
    ])
    m.table(
        ["Outcome", "What it does"],
        [
            ["Reinstated", "Back to work. The account unlocks and the status returns to Active, the same day."],
            ["Withdrawn", "It should not have happened. The same effect on the account, a different word in the record."],
            ["Terminated", "An exit. The roster status becomes Terminated and the offboarding flow picks it up."],
        ],
        widths=[1.3, 5.2],
    )
    m.note(
        "Rehiring somebody who is suspended is refused. A suspension is not a "
        "departure. Lift it instead.",
        "warn",
    )

    m.h2("4.6 Change requests")
    m.p(
        "An employee cannot silently change their own bank account or "
        "citizenship number. Those edits become a request that HR sees under "
        "People, Change requests, with the old value beside the new one and the "
        "reason they gave. Payroll is downstream of that queue, which is why it "
        "exists."
    )

    m.h2("4.7 Work history")
    m.p(
        "Two kinds, kept in one list because they carry the same facts, and shown "
        "as two sections because they mean different things:"
    )
    m.bullets([
        "Previous employment: somewhere else, before joining. Self-declared, and marked unverified until HR checks it against a document.",
        "Held here: a post held inside the group. A fact the system wrote itself.",
    ])
    m.p(
        "An employee maintains their own entries from their profile. HR "
        "maintains anybody's, and the verified tick can only be set by HR. A "
        "claim that verifies itself is not a check."
    )


def memoranda(m: Manual) -> None:
    m.h1("5. Memoranda")
    m.p(
        "A memorandum is a proposal that goes up a chain: the person with the "
        "idea writes it, named people recommend it in an order, and one person "
        "approves or refuses it. This is the largest workflow in the system and "
        "the one with the most rules, so this chapter sets out the rules first "
        "and the buttons second."
    )

    m.h2("5.1 The desk")
    m.figure("fig-memoranda", "Figure 5.1 — The memorandum desk")
    m.table(
        ["#", "What it is", "What it does"],
        [
            ["1", "New memorandum", "Raise one. Available to everybody with an employee record."],
            ["2", "Actions", "The vocabulary table. Everybody can read it; only the owner or an HR admin can change it."],
            ["3", "Search", "Number, subject or content, across everything you can see. Switches the page to a flat result list."],
            ["4", "Needs you", "Waiting on you right now. Highlighted, and first, because a memorandum sitting on a desk nobody knows about is the failure of the paper system this replaces."],
            ["5", "Raised by you", "Yours, wherever they have got to."],
            ["6", "You have handled", "History. Where you go to find the note you signed in Poush."],
        ],
        widths=[0.35, 1.5, 4.65],
    )

    m.h2("5.2 The letter")
    m.p(
        "A memorandum opens as the sheet of paper it replaces: a letterhead with "
        "the company and its seat, then Ref and Date, then To and From, then the "
        "Subject, then the body, and a signature at the foot. There are no tabs "
        "and nothing to switch between."
    )
    m.rich(
        ("Where a field can still be changed, the control sits where its value "
         "would be printed", True),
        (". The company picker is the letterhead. The subject box is the subject "
         "line. The editor is the body. Filling one in is filling in the letter, "
         "which is the point: there is nothing to learn beyond where to click.",
         False),
    )
    m.table(
        ["Line", "What goes there", "When it can be changed"],
        [
            ["Ref", "The memorandum number", "Issued by the system at submission. Never typed."],
            ["Date", "The date it is raised", "While it is a draft. It must be today when you submit."],
            ["To", "The approver", "Until it reaches them. Chosen under Who signs it."],
            ["From", "You, and the post you hold", "Never — it is whoever raised it."],
            ["Subject", "One line, what this is about", "While it is a draft."],
            ["Body", "The proposal itself", "While anybody in the chain still has it, by the initiator. The one field that survives submission."],
        ],
        widths=[0.7, 2.4, 3.4],
    )
    m.note(
        "The status and who is holding it sit on a tab clipped to the top-right "
        "corner, outside the sheet. A real memorandum does not carry a chip "
        "reading “in progress” — that is something the system knows about the "
        "document, not something the document says.",
        "note",
    )

    m.h2("5.3 Raising one")
    m.figure("fig-memorandum-new", "Figure 5.2 — A new memorandum")
    m.h3("Step by step")
    m.steps([
        "Click Memorandum in the sidebar, under Workplace.",
        "Click New memorandum, at the top right. A blank letter opens.",
        "Company — click the letterhead at the top of the page and choose one of the four entities. Its code becomes part of the memorandum number, so it cannot be changed after submission.",
        "Date — on the right of the Ref line, already filled in with today.",
        "Subject — click the subject line and type one line saying what this is about.",
        "Body — click into the page and write. The toolbar gives bold, italic, underline, strike-through, font size, bulleted and numbered lists, alignment, and undo.",
        "Scroll below the page to Who signs it.",
        "Recommenders — add people in the order they should see it. The order you pick them is the order they get it.",
        "Approver — one person, beside the recommenders. They become the To of the letter.",
        "Click Save draft. The dialog reopens on the saved draft.",
        "Attachments — now use Choose a file, add a caption, and click Attach. Repeat for as many as you need.",
        "Click Submit.",
    ])
    m.figure("fig-memorandum-chain", "Figure 5.3 — Below the page: attachments, who signs it, and the history")
    m.note(
        "The dialog reopens rather than closing when you save a new draft. "
        "Attachments hang off a memorandum, so one has to exist before there is "
        "anything to attach them to.",
        "note",
    )

    m.h3("The number")
    m.rich(
        ("On submission the memorandum is given a number in the form ", False),
        ("yyyy-mm-dd-CODE-serial", True),
        (", for example 2026-09-02-VLUCL-0001. It is minted at submission and "
         "never at creation, so a draft that is abandoned does not consume a "
         "number out of the company's register. The serial runs per company.", False),
    )

    m.h2("5.4 Handling one")
    m.figure("fig-memorandum-open", "Figure 5.4 — A memorandum in flight")
    m.p(
        "When a memorandum reaches you it appears under Needs you. Open it and "
        "the action bar offers exactly what you may do at that point."
    )
    m.table(
        ["Action", "What happens"],
        [
            ["Send it on", "Choose a word from the dropdown — Recommended, Noted, Verified, whatever the owner has configured — add a comment if you want, and it moves to the next person."],
            ["Send it back", "Choose who to send it back to, and why. The choices are the initiator and anybody earlier in the chain who has already handled it."],
            ["Approve", "Approver only, and only at the end. The memorandum closes."],
            ["Reject", "Approver only. The memorandum closes."],
            ["Comment", "Available to anybody who can see it, at any point, without moving it."],
        ],
        widths=[1.2, 5.3],
    )

    m.h2("5.5 Comments, mentions and files")
    m.p(
        "Anybody who can see a memorandum can comment on it — including "
        "somebody three steps down the chain who spots a problem before it "
        "reaches them. A comment can carry two extra things."
    )
    m.note(
        "There is one comment box per person, and where it is depends on whose "
        "turn it is. Whoever is holding the memorandum comments in the action "
        "panel, where the note travels with the decision. Everybody else uses "
        "the box below the history. Two boxes that look the same, one of which "
        "silently attaches to whichever button you press next, is a choice "
        "nobody should have to make.",
        "note",
    )
    m.h3("Naming somebody")
    m.steps([
        "Type the comment.",
        "Click the Notify box and pick the people who should see it.",
        "Click Comment.",
    ])
    m.p(
        "They are told, and they can then open and read the memorandum. It "
        "grants reading and nothing else: acting on a memorandum still means "
        "being the person holding it. This exists because the chain is chosen "
        "for who must decide, and the people who know the answer are usually "
        "not in it."
    )
    m.h3("Attaching a file to a comment")
    m.steps([
        "Click Attach and pick one or more files. They appear as chips under the box.",
        "Remove any you did not mean by clicking the x on the chip.",
        "Click Comment.",
    ])
    m.note(
        "A file on a comment is not one of the memorandum's annexes and does not "
        "appear in its attachment list. The annexes are part of the proposal and "
        "freeze at submission; a file on a comment is a reply — usually the "
        "document somebody was sent back to fetch.",
        "note",
    )

    m.h2("5.6 The rules")
    m.p(
        "These are enforced by the system, not by convention. Each one refuses "
        "with a sentence saying why."
    )
    m.table(
        ["Rule", "Why"],
        [
            ["A memorandum cannot be backdated", "It is dated the day it is raised. One dated last week and submitted now is either a mistake or a register being rewritten."],
            ["After submission only the content can change", "Not the subject, not the date, not the company. Sending a memorandum back is what editing the rest is for."],
            ["A recommender who has acted cannot be removed", "Their comment is part of the record, and the chain is what it is attached to. Removing them would leave a signature on a document that no longer says they were involved."],
            ["Neither can the person currently holding it", "Ask them to send it back first."],
            ["Somebody not yet reached can be removed or reordered", "This is the normal case: it comes back saying finance should see it too, and the initiator adds them."],
            ["The approver can be changed until it reaches them", "At that point they are reading it, and swapping the person mid-read takes a document off somebody's desk with no explanation."],
            ["Only the person holding it can act on it", "Anybody else is told This memorandum is not with you."],
            ["After a decision, nothing changes", "Not the text, not the chain, not the comments. It is a record."],
        ],
        widths=[2.3, 4.2],
    )

    m.h2("5.7 The cycle")
    m.p(
        "Returning is not a dead end and not a reset. The memorandum is at a "
        "position in its chain; sending it back moves that position backwards, "
        "and it then comes forward again through the same people. The loop can "
        "run any number of times."
    )
    m.p("A worked example, with four recommenders A, B, C and D:")
    m.steps([
        "The initiator submits. It lands on A.",
        "A sends it on. It lands on B.",
        "B sends it back to the initiator, asking for a survey.",
        "The initiator attaches the survey and clicks Send forward again.",
        "It lands on A once more — the point it was returned from — and climbs from there.",
        "A, B, C and D each send it on. It reaches the approver.",
        "The approver approves it. It closes, and nothing on it can be changed again.",
    ])
    m.note(
        "At step 3 the initiator could not have removed A from the chain: A had "
        "already acted. They could still have removed C or D, who had not.",
        "warn",
    )

    m.h2("5.8 The history")
    m.figure("fig-memorandum-history", "Figure 5.5 — The full history")
    m.p(
        "Every memorandum carries a complete log: who did what, in what words, "
        "when, and to whom it was sent back. It is append-only. A correction is "
        "a new entry, because a log that could be rewritten would make the "
        "record worth less than the paper version it replaced."
    )
    m.p("The words in the log are frozen when the entry is written. Renaming an action later does not rewrite history.")

    m.h2("5.9 The action vocabulary")
    m.figure("fig-memorandum-actions", "Figure 5.6 — The configurable action table")
    m.rich(
        ("Every organisation argues over its own words — recommended, noted, "
         "reviewed, verified, supported — so they are configuration rather than "
         "code. The system reads only the ", False),
        ("effect", True),
        (": a word either sends the memorandum on, or sends it back. The wording "
         "is yours.", False),
    )
    m.h3("Step by step: adding a word")
    m.steps([
        "Go to Memorandum and click Actions, or Settings then Memorandum actions.",
        "Click Add an action.",
        "Word — what appears in the log. For example, Endorsed.",
        "Code — a short identifier, upper case.",
        "Effect — choose Send it on or Send it back. This is the only part the system reads.",
        "Order — where it sits in the dropdown.",
        "Meaning — what choosing it says. Shown to whoever is picking.",
        "Also offered to the approver — usually off. Recommended is not something an approver says, and offering it there leaves a memorandum neither approved nor refused with nowhere to go.",
        "Click Save.",
    ])
    m.note(
        "A word that has been used cannot be deleted. Deactivate it instead: it "
        "leaves every dropdown and the history stays readable.",
        "warn",
    )


def field_visits(m: Manual) -> None:
    m.h1("6. Field visits")
    m.figure("fig-field-visits", "Figure 6.1 — Field visits")
    m.rich(
        ("A field visit is a travel order: approved ", False),
        ("before", True),
        (" departure, reported on return.", False),
    )

    m.h2("6.1 Why this is not a timesheet")
    m.p(
        "The question was asked, and the answer is worth recording. A time entry "
        "is a scalar: hours, against a project, on a day, recorded afterwards "
        "and approved in bulk. A visit is a journey — it has a destination, a "
        "purpose, companions, a cost, an order approved before anybody sets off, "
        "and a report afterwards which is the thing the visit existed to "
        "produce. Folding one into the other means either a time entry that can "
        "be approved in advance, or four more empty columns on the busiest table "
        "in the system."
    )
    m.p("They are joined at three points instead, and each is worth knowing:")
    m.table(
        ["Seam", "What it does"],
        [
            ["Attendance", "An approved visit keeps the traveller off the absentee list. Somebody at site for a week has no clock-in for five days, and without this the nightly sweep records five absences, which feed unpaid days and cut their pay. The company would be docking the pay of the person it sent."],
            ["Timesheets", "A completed visit can generate one time entry per day against its project. Run it twice and the second run adds nothing."],
            ["Expenses", "A visit links to one expense claim, so the advance has something to sit on."],
        ],
        widths=[1.1, 5.4],
    )

    m.h2("6.2 Raising and running a visit")
    m.figure("fig-field-visit-new", "Figure 6.2 — A new field visit")
    m.h3("Step by step")
    m.steps([
        "Click Field visits in the sidebar, under Time and attendance. It sits next to Timesheets on purpose: they are the two things people confuse.",
        "Click New visit.",
        "Fill in what the visit is for, the destination and the district.",
        "Purpose — click the dropdown: inspection, construction supervision, survey, maintenance, emergency response, meeting, community, audit, training or other.",
        "Company, dates, project, approver, transport and estimated cost.",
        "Click Save as draft. The visit opens.",
        "Click Send for approval.",
        "The approver opens it and clicks Approve or Reject, with a note.",
        "Who else went — add companions under the report. A name is all that is required; picking a staff member fills it in. Somebody who does not work here is recorded by name, with an organisation and their part.",
        "On return, the traveller writes the report and clicks Complete with report.",
        "Optionally, click Write timesheet lines to turn the visit into time entries.",
    ])
    m.note(
        "A visit cannot be closed without a report. A visit with no findings is "
        "a cost with no output, and the report is the only part of the record "
        "anybody reads a year later.",
        "warn",
    )
    m.p(
        "Only approved and completed visits count for attendance. A request "
        "nobody has signed is a plan, not a reason to be away — otherwise anybody "
        "could keep themselves off the absentee list by writing one."
    )


def payroll(m: Manual) -> None:
    m.h1("7. Payroll")
    m.figure("fig-payroll-run", "Figure 7.1 — A payroll run")

    m.h2("7.1 The stages of a run")
    m.table(
        ["Stage", "What it means", "What you can do"],
        [
            ["Draft", "Created, nothing computed", "Run Payroll"],
            ["Processing", "Computing in the background", "Wait. The list updates itself."],
            ["Completed", "Payslips computed, still editable", "Review, edit a payslip, Finalize"],
            ["Locked", "Finalised. Figures immutable", "Mark paid, download, build payment batches"],
        ],
        widths=[0.9, 2.4, 3.2],
    )

    m.h2("7.2 Step by step: running payroll")
    m.steps([
        "Click Payroll in the sidebar, under Finance.",
        "Open the run for the period, or create one.",
        "Click Run Payroll. The status becomes Processing while payslips compute in the background.",
        "When it reads Completed, review the payslips. Filter by Draft, Finalised or Paid using the chips.",
        "Correct anything that needs it — open a payslip and edit its line items, or click Recompute to discard edits and re-prefill from the salary structure.",
        "Click Finalize. This locks the period.",
        "Click Mark all paid, or mark payslips individually, once the money has actually gone.",
        "Download payslips as PDF, or export a bank file for the corporate portal.",
    ])

    m.h2("7.3 Why Finalize refuses")
    m.p(
        "Finalising locks the period and there is no undo, so it is guarded. "
        "Both refusals name exactly what to fix and link to the page that fixes "
        "it."
    )
    m.table(
        ["Refusal", "What it means", "What to do"],
        [
            ["This run has unresolved errors", "One or more payslips failed to compute. The message names whose.", "Fix the underlying record, resolve the error, run again."],
            ["N statutory figures have not been checked", "Rates and tax bands ship marked unchecked until somebody confirms them against the Finance Act.", "Open Payroll, Statutory rates and verify each row."],
        ],
        widths=[1.8, 2.5, 2.2],
    )
    m.figure("fig-statutory-rates", "Figure 7.2 — Statutory rates, and the verified flag")
    m.note(
        "There is no override on either guard, deliberately. An override would "
        "be taken every time by whoever is in a hurry, and on payroll day that "
        "is everybody. Verifying is one click per row.",
        "stop",
    )

    m.h2("7.4 How a payslip is computed")
    m.p("In order, and each step reads only what the step before it produced:")
    m.steps([
        "Read the employee's salary structure as it stood on the period start date. A structure is versioned, so a raise in the middle of a year does not restate last year.",
        "Compute each earning component. Flat amounts are taken as they are; percentage components are taken of whatever they name; formula components are evaluated.",
        "Sum the earnings into gross.",
        "Prorate for the days actually employed, if the company has proration on and the person did not work the whole period.",
        "Compute statutory deductions: SSF or Provident Fund, but never both, and gratuity where it applies.",
        "Compute income tax from the fiscal year's slabs for that taxpayer type, allowing the retirement-relief ceilings and the female rebate where applicable.",
        "Apply loan repayments due this period.",
        "Apply any manual adjustments HR has entered.",
        "Net pay is gross less the sum of deductions.",
    ])
    m.p(
        "Every line on a payslip says where its number came from, which is the "
        "point: a payslip that shows only totals cannot be argued with, and "
        "payroll queries are always about one line."
    )


def other_modules(m: Manual) -> None:
    m.h1("8. The other modules")

    m.h2("8.1 Companies")
    m.figure("fig-companies", "Figure 8.1 — Companies")
    m.p(
        "The four entities, their registration and licence numbers, installed "
        "capacity, river, project stage and seat. A company cannot be its own "
        "parent, directly or through a loop. Companies are protected from "
        "deletion while anybody is employed by them."
    )

    m.h2("8.2 Events")
    m.figure("fig-events", "Figure 8.2 — Events")
    m.p(
        "Public hearings, inaugurations, board meetings, community programmes. "
        "Shown as a timeline of past and upcoming rather than a calendar grid, "
        "because an event is remembered by what it was, not by which Tuesday it "
        "fell on."
    )
    m.rich(
        ("Stakeholders take a ", False),
        ("name first and an employee second", True),
        (". Half the people at a public hearing are not staff — a ward chair, a "
         "contractor's foreman, a ministry official — and a list that could only "
         "name employees would record four colleagues and silently drop eleven "
         "others. Picking an employee fills the name in for you.", False),
    )

    m.h2("8.3 Attendance and leave")
    m.figure("fig-attendance", "Figure 8.3 — Attendance")
    m.bullets([
        "Logs, corrections, shifts and shift assignments",
        "A monthly calendar per person, in Bikram Sambat or Gregorian",
        "Lateness measured against the assigned shift, not a fixed hour",
        "Absence is swept nightly, and skips non-working days, holidays and anybody on an approved field visit",
    ])
    m.figure("fig-leave", "Figure 8.4 — Leave")
    m.p(
        "Leave types with their own accrual, balances per person, requests that "
        "go to the manager, and a calendar showing who is away. Leave and "
        "attendance read the same working week and holiday list, so they cannot "
        "disagree about whether a day counts."
    )

    m.h2("8.4 Expenses, and budgets")
    m.figure("fig-expenses", "Figure 8.5 — Expense claims")
    m.figure("fig-expense-budgets", "Figure 8.6 — Budgets and caps")
    m.p(
        "Claims with receipts, an approval step, and reimbursement. Budgets sit "
        "over the top: a cap per category, per period, optionally narrowed to a "
        "department or a single person. The most specific budget that matches "
        "wins. Pending claims count against the cap, because a budget that only "
        "counted approved spending would let a queue of claims sail past it."
    )

    m.h2("8.5 Assets")
    m.figure("fig-assets", "Figure 8.7 — Assets")
    m.p(
        "Equipment, who holds it, and its history. The history is the part worth "
        "knowing about: every assignment, return, maintenance and retirement is "
        "recorded with dates, so the question who had this laptop in Ashad has "
        "an answer."
    )

    m.h2("8.6 Training")
    m.figure("fig-training", "Figure 8.8 — Training")
    m.figure("fig-training-program", "Figure 8.9 — A programme, and its sessions")
    m.rich(
        ("The catalogue is open to ", False),
        ("everybody", True),
        (". It is the one HR module an ordinary employee is meant to open, and "
         "there are two ways into a session:", False),
    )
    m.h3("An employee asks")
    m.steps([
        "Open Training and click a programme.",
        "Find the session and click Request to join.",
        "HR opens the session's Roster and clicks Approve or Decline.",
    ])
    m.h3("HR invites")
    m.steps([
        "Open the programme and click Roster on the session.",
        "Add participants directly. No request is needed.",
    ])
    m.p("On completion, a certificate is issued and downloads as a PDF.")

    m.h2("8.7 Helpdesk, recruitment, timesheets")
    m.figure("fig-helpdesk", "Figure 8.10 — Internal helpdesk")
    m.bullets([
        "Helpdesk — internal tickets, with a selectable handler so a request reaches the right desk rather than a shared queue",
        "Recruitment — jobs, candidates and a pipeline board; a candidate becomes an employee in one step",
        "Timesheets — hours against projects, with approval",
    ])
    m.figure("fig-recruitment", "Figure 8.11 — Recruitment")
    m.figure("fig-timesheets", "Figure 8.12 — Timesheets")

    m.h2("8.8 The calendar")
    m.figure("fig-calendar", "Figure 8.13 — The company calendar")
    m.p(
        "Opens in Bikram Sambat, with a toggle to Gregorian. Holidays, meetings, "
        "events and announcements on one grid. The Gregorian view supports "
        "drag-to-reschedule; the Bikram Sambat grid does not, because "
        "reimplementing that interaction on a second grid would be two "
        "behaviours to keep in step."
    )

    m.h2("8.9 Settings")
    m.figure("fig-settings", "Figure 8.14 — Settings")
    m.table(
        ["Page", "What it holds"],
        [
            ["Company", "Name, working week, fiscal year, calendar"],
            ["Departments and job titles", "The organisation's own vocabulary"],
            ["Holidays", "The year's list, which leave and attendance both read"],
            ["Salary components", "Earnings and deductions, and how each is computed"],
            ["Tax slabs", "By fiscal year and taxpayer type"],
            ["Statutory rates", "SSF, PF, gratuity, ceilings, and the verified flag"],
            ["Memorandum actions", "The approval vocabulary"],
            ["Email", "Outgoing mail. Until this is set, nobody you add can sign in, because their temporary password is emailed to them."],
            ["Attendance", "Shift rules, lateness thresholds"],
            ["Notifications and reminders", "What the system sends, and when"],
        ],
        widths=[1.6, 4.9],
    )


def by_user(m: Manual) -> None:
    m.h1("9. What each person does")
    m.p(
        "The same system, read from four desks. Each list is what that person "
        "would actually do in a week, in the order they would do it."
    )

    m.h2("9.1 An employee")
    m.bullets([
        "Sign in, and choose a password on first use",
        "Open My workspace for what is waiting on them",
        "Request leave, and see the balance it comes out of",
        "Clock in and out, and query a wrong attendance entry",
        "See payslips, and download any of them as PDF",
        "Ask to change their own bank account or address — it becomes a request HR decides",
        "Submit an expense claim with receipts",
        "Browse Training, and ask for a seat on a session",
        "Raise a memorandum, and act on any that name them",
        "Ask for a field visit, and write the report on return",
        "Fill in a timesheet",
        "Raise a helpdesk ticket",
    ])
    m.p("An employee cannot see anybody else's private details, payslips, bank account or conduct record.")

    m.h2("9.2 A manager")
    m.p("Everything an employee does, plus, for the people whose Reports to field names them:")
    m.bullets([
        "Approve or refuse their leave",
        "See their attendance",
        "Approve their timesheets and field visits",
        "Appear above them on the organisation chart",
    ])
    m.note(
        "Manager is not a role. It is the Reports to field, and it grants reach "
        "over those people without granting any capability of its own.",
        "note",
    )

    m.h2("9.3 An HR officer")
    m.bullets([
        "Open and edit any employee record",
        "Maintain conduct records, work history, contacts and dependants",
        "Suspend and lift suspensions",
        "Correct attendance, manage shifts, approve leave",
        "Run assets, training, helpdesk, timesheets and remote work",
        "Read reports and the dashboard",
    ])
    m.p("And cannot, without a grant from the owner:")
    m.bullets([
        "Add a new employee, or anything else",
        "Delete anything at all",
        "Open payroll, or see anybody's salary",
        "Change settings — components, slabs, the memorandum vocabulary",
        "Read the company mailbox",
    ])

    m.h2("9.4 An HR admin")
    m.bullets([
        "Everything an officer does, plus creating and deleting",
        "Run payroll end to end, and finalise it",
        "Verify statutory rates and tax slabs",
        "Define salary components and the memorandum vocabulary",
        "Manage companies, departments and job titles",
        "Promote and demote officers and employees",
    ])
    m.p("And cannot: appoint or demote another HR admin. That stays with the owner.")

    m.h2("9.5 The owner")
    m.bullets([
        "Everything",
        "Appoint HR admins — the only role that can",
        "Reach the Django administration site at /admin/",
    ])
    m.note(
        "In this installation the owner account has no employee record attached. "
        "That is deliberate — the owner is the account the system was installed "
        "under, not a member of staff — and it has one consequence worth knowing: "
        "the owner cannot raise a memorandum, request leave, or appear in an "
        "approval chain, because all three belong to a person. Use an HR admin "
        "account for those.",
        "warn",
    )


def algorithms(m: Manual) -> None:
    m.h1("10. How the important things are decided")
    m.p(
        "Four processes where the order of steps matters, set out in full. "
        "These are the rules the software runs, not a summary of them."
    )

    m.h2("10.1 Signing in")
    m.steps([
        "The username and password are posted to the server. The browser never holds the resulting token — it is kept in a cookie the page's own scripts cannot read.",
        "The account is looked up. If it does not exist, or the password is wrong, the answer is that no active account matches. The message does not distinguish the two, because doing so tells an attacker which usernames are real.",
        "If the account is locked, the reason is checked. A lock caused by a suspension answers with the suspension and its end date, rather than pretending the password was wrong.",
        "If two-factor authentication is on, a code is asked for and the sign-in pauses there.",
        "On success a short-lived access token and a longer refresh token are issued. The access token expires quickly; the refresh happens silently.",
        "Permissions are read on every request, never baked into the token. A capability revoked this morning is gone this morning — a permission carried in a token would survive until it expired, which is a revocation that does not revoke.",
    ])

    m.h2("10.2 A memorandum, from raised to closed")
    m.steps([
        "Created. Status draft, no number. The initiator is whoever is signed in — it is not a field on the form.",
        "The content is cleaned as it is stored. Formatting is kept; anything that could run is removed. This happens on the server, so it holds however the content arrived.",
        "Recommenders are stored with an order. The approver is stored separately.",
        "Submitted. The date is checked against today, and refused if it differs.",
        "A serial is taken for that company, under a lock so two people submitting at the same instant cannot take the same one. The number is composed and stored.",
        "The cursor is set to position 0 and the memorandum lands on the first recommender. Status becomes in progress, stage becomes recommend.",
        "Sent on: the cursor advances by one. If that is past the last recommender, the stage becomes approve and it lands on the approver.",
        "Sent back: the cursor moves to the chosen person's position, or to the initiator. There is no separate returned state, which is why the loop can run any number of times with nothing to unwind.",
        "Every move writes a line in the log, with the actor, the word used and the time, all frozen at that moment.",
        "Approved or rejected: the stage becomes closed and the memorandum is locked. Every subsequent write is refused.",
    ])
    m.p(
        "Whether somebody has acted is read from the log rather than from a flag "
        "on the chain. A return would have to clear such a flag, and the question "
        "the rules need answered is whether they have ever acted, not whether "
        "they have since the last loop."
    )

    m.h2("10.3 The nightly absence sweep")
    m.steps([
        "The date is checked against the company's working week. A Saturday marks nobody.",
        "It is checked against the holiday list. A holiday marks nobody.",
        "Everybody with an active shift assignment on that date is gathered — that is what scheduled to work means.",
        "Anybody who already has an attendance log for the date is removed from the list.",
        "For each person left, an approved or completed field visit covering the date is looked for. If there is one, a Present log is written naming the destination, and they are skipped.",
        "Everybody still remaining is marked Absent.",
    ])
    m.note(
        "Steps 1 and 2 are not politeness. Absences feed unpaid days, which "
        "scale pay directly. Without the working-week check a Monday-to-Friday "
        "company running this daily would accrue about nine absences a month per "
        "person and cut every salary by roughly a third.",
        "stop",
    )

    m.h2("10.4 Suspension, and the account lock")
    m.steps([
        "A suspension record is created with a start date, an optional end date and a reason.",
        "In one transaction: the record is marked active, the employment status becomes Suspended, and the account is deactivated.",
        "Because the account is deactivated, every request that account makes is refused from that moment. This includes the token it already holds — the check runs per request, not at sign-in.",
        "A notice is emailed, to the personal address if there is one, otherwise the office address. If mail fails, the suspension stands: the lock is the record, the notice is not.",
        "On lifting, an outcome is required. Reinstated and Withdrawn restore the status and reactivate the account. Terminated sets the status to Terminated and leaves the account locked.",
        "A suspension with an end date that has passed is closed by the same nightly sweep, so a fixed-period suspension ends without anybody remembering to end it.",
    ])


def limitations(m: Manual) -> None:
    m.h1("11. Limitations")
    m.p(
        "What the system does not do, or does only partly. Listed so that "
        "nobody spends an afternoon looking for a feature that is not there."
    )

    m.h2("11.1 Not built")
    m.table(
        ["Area", "Limitation"],
        [
            ["Deleting", "Only unsent drafts can be deleted, and only by whoever raised them. A submitted memorandum, an approved field visit and anything else with a trail behind it are never deletable — by design, not omission."],
            ["Single sign-on", "Not implemented. There is no Google or Microsoft sign-in, and on a single-company installation there is no second identity provider to federate with."],
            ["Self-service sign-up", "There is none, deliberately. Accounts are created from inside the product."],
            ["Public website", "Removed. This build has no marketing site or careers portal."],
            ["Mobile application", "None. The interface adapts to a phone browser but there is no installable app."],
            ["Biometric devices", "The data model supports device ingestion; no device is connected in this installation."],
        ],
        widths=[1.6, 4.9],
    )

    m.h2("11.2 Deliberate restrictions")
    m.p("These look like limitations and are decisions:")
    m.bullets([
        "Payroll cannot be finalised on unverified statutory figures, and there is no override.",
        "A memorandum cannot be backdated, and cannot be edited after approval.",
        "A recommender who has acted cannot be removed from a chain.",
        "An HR officer cannot create or delete anything.",
        "Suspended cannot be set from the status dropdown.",
        "A field visit cannot be closed without a report.",
        "Attachments on a memorandum freeze at submission.",
    ])

    m.h2("11.3 Operational notes")
    m.table(
        ["Area", "Note"],
        [
            ["Testing quality", "The server has 1,341 automated tests. The interface has none, and is verified by hand — so a fault in the browser layer is likelier to reach a user than one in the server."],
            ["Mail", "Points at a local inbox in this installation. Configure real SMTP under Settings, Email before anything reaches a real address."],
            ["Source changes", "The frontend container serves a built copy. A code change needs the container restarted before it appears."],
            ["Browser extensions", "A password manager may inject attributes into the sign-in form and produce a console warning about a hydration mismatch. It comes from the extension, not the system."],
        ],
        widths=[1.4, 5.1],
    )


def appendix(m: Manual) -> None:
    m.h1("12. Appendix — demonstration accounts")
    m.note(
        "These belong to the seeded demonstration database. They are fictional "
        "people, and example.com is reserved by the IETF precisely so it can "
        "never belong to anybody. Remove them before this installation carries "
        "real staff.",
        "stop",
    )
    m.rich(("The password for every account below is ", False), ("TestPass123!", True), ("", False))
    m.table(
        ["Role", "Username", "Employee", "Use it to see"],
        [
            ["Owner", "owner", "—", "Everything, plus the Django administration site"],
            ["HR admin", "sushma.ghimire", "EMP-0022", "Creating, deleting, payroll, settings"],
            ["HR officer", "tenzing.neupane", "EMP-0079", "Editing without creating or deleting"],
            ["HR officer", "rachana.baral", "EMP-0094", "A second officer, for two people on one record"],
            ["Employee", "rajendra.poudel", "EMP-0001", "The employee view, and a memorandum chain"],
            ["Employee", "bikash.shrestha", "EMP-0002", "A recommender on a seeded memorandum"],
            ["Employee", "kabita.thapa", "EMP-0006", "An approver on a seeded memorandum"],
            ["Suspended", "sarita.gurung", "—", "The sign-in lock. This account cannot sign in."],
        ],
        widths=[0.95, 1.5, 0.95, 3.1],
    )
    m.p("There are 96 further employee accounts, all named firstname.lastname.")

    m.h2("12.1 What the demonstration data contains")
    m.table(
        ["Records", "Count"],
        [
            ["Employees", "105"],
            ["Companies", "4"],
            ["Attendance logs", "1,892"],
            ["Payroll runs / payslips", "6 / 72"],
            ["Memoranda", "7, in every state the chain can reach"],
            ["Field visits", "5, one covering today"],
            ["Events", "6, past and upcoming"],
            ["Leave requests", "51"],
            ["Expense claims / budgets", "22 / 4"],
            ["Assets / asset history", "14 / 18"],
            ["Training programmes / enrolments", "4 / 28"],
        ],
        widths=[3.0, 3.5],
    )

    m.h2("12.2 Resetting the demonstration data")
    m.p("From a terminal on the machine running the system:")
    m.table(
        ["Step", "Command"],
        [
            ["Stop and drop the database", "docker compose down -v"],
            ["Start again", "docker compose up -d"],
            ["Reseed", "docker exec hrms-hydro-backend python manage.py seed_hydro --owner-password 'TestPass123!'"],
        ],
        widths=[1.9, 4.6],
    )
    m.note(
        "The -v flag removes the volume, which is what makes it a genuine clean "
        "slate. Without it the database survives and the seed layers on top of "
        "whatever is there.",
        "warn",
    )


def build(figures: Path, out: Path) -> Path:
    m = Manual(figures)
    cover(m)
    contents(m)
    introduction(m)
    access(m)
    roles(m)
    employees(m)
    memoranda(m)
    field_visits(m)
    payroll(m)
    other_modules(m)
    by_user(m)
    algorithms(m)
    limitations(m)
    appendix(m)
    return m.save(out)


if __name__ == "__main__":
    figures = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs/figures")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("Xenex-HRMS-User-Manual.docx")
    path = build(figures, out)
    print(f"written -> {path}  ({path.stat().st_size / 1_048_576:.1f} MB)")
